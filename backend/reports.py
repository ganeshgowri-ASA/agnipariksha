"""Report generation for Agnipariksha test sessions.

Produces branded Word (.docx) and PDF reports for a given session,
including Pass/Fail verdict, IEC reference, a matplotlib chart of the
measurement trace, and a raw-data table.

Reports are written to REPORTS_DIR (default ./reports/) and their
path is recorded in the `reports` table.
"""
from __future__ import annotations

import io
import json
import os
from datetime import datetime
from typing import List, Optional

import matplotlib

matplotlib.use("Agg")  # headless backend, required on servers
import matplotlib.pyplot as plt

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

from reportlab.lib import colors
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import cm
from reportlab.platypus import (
    Image as RLImage,
    Paragraph,
    SimpleDocTemplate,
    Spacer,
    Table,
    TableStyle,
)

from sqlalchemy.orm import Session

from db import Measurement, Report, TestSession


BRAND_NAME = "Shreshtata Power Supplies"
PRODUCT_NAME = "Agnipariksha"
BRAND_TAGLINE = "PV Module Reliability Test Station"
BRAND_PRIMARY = "#0B6E4F"   # deep green
BRAND_ACCENT = "#FFB100"

REPORTS_DIR = os.getenv("AGNI_REPORTS_DIR", os.path.join(os.path.dirname(__file__), "reports"))
os.makedirs(REPORTS_DIR, exist_ok=True)


# ----- shared helpers --------------------------------------------------------

def _load(db: Session, session_id: str) -> tuple[TestSession, List[Measurement]]:
    s = db.get(TestSession, session_id)
    if not s:
        raise ValueError(f"session {session_id} not found")
    rows = (
        db.query(Measurement)
        .filter(Measurement.session_id == session_id)
        .order_by(Measurement.ts.asc())
        .all()
    )
    return s, rows


def _make_chart_png(rows: List[Measurement]) -> bytes:
    """Render voltage/current/power vs time as a PNG (bytes)."""
    fig, ax1 = plt.subplots(figsize=(7.5, 3.5), dpi=120)
    if not rows:
        ax1.text(0.5, 0.5, "No measurements recorded", ha="center", va="center")
        ax1.set_axis_off()
    else:
        t0 = rows[0].ts
        ts = [(r.ts - t0).total_seconds() for r in rows]
        vs = [r.v or 0.0 for r in rows]
        is_ = [r.i or 0.0 for r in rows]
        ps = [r.p or 0.0 for r in rows]

        ax1.plot(ts, vs, color=BRAND_PRIMARY, label="V (V)", linewidth=1.2)
        ax1.plot(ts, is_, color=BRAND_ACCENT, label="I (A)", linewidth=1.2)
        ax1.set_xlabel("Elapsed (s)")
        ax1.set_ylabel("Voltage / Current")
        ax1.grid(True, alpha=0.3)

        ax2 = ax1.twinx()
        ax2.plot(ts, ps, color="#444", label="P (W)", linewidth=0.9, linestyle="--")
        ax2.set_ylabel("Power (W)")

        lines1, labels1 = ax1.get_legend_handles_labels()
        lines2, labels2 = ax2.get_legend_handles_labels()
        ax1.legend(lines1 + lines2, labels1 + labels2, loc="upper right", fontsize=8)

    fig.tight_layout()
    buf = io.BytesIO()
    fig.savefig(buf, format="png")
    plt.close(fig)
    buf.seek(0)
    return buf.getvalue()


def _verdict_and_color(result: Optional[dict]) -> tuple[str, str]:
    if not result:
        return "INCONCLUSIVE", "#666666"
    v = (result.get("verdict") or "INCONCLUSIVE").upper()
    color = {"PASS": "#0B6E4F", "FAIL": "#B00020"}.get(v, "#666666")
    return v, color


# ----- Word report -----------------------------------------------------------

def generate_word(db: Session, session_id: str) -> str:
    s, rows = _load(db, session_id)
    chart_png = _make_chart_png(rows)
    verdict, _ = _verdict_and_color(json.loads(s.result_json) if s.result_json else None)

    doc = Document()

    # Header / brand block
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run(f"{BRAND_NAME} — {PRODUCT_NAME}")
    run.bold = True
    run.font.size = Pt(20)
    run.font.color.rgb = RGBColor.from_string(BRAND_PRIMARY.lstrip("#"))

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub_run = sub.add_run(BRAND_TAGLINE)
    sub_run.italic = True
    sub_run.font.size = Pt(11)

    doc.add_paragraph()  # spacer

    # Metadata table
    meta = doc.add_table(rows=0, cols=2)
    meta.style = "Light Grid Accent 1"

    def _row(k: str, v: str) -> None:
        r = meta.add_row().cells
        r[0].text = k
        r[1].text = v

    _row("Session ID", s.id)
    _row("Test", s.test_id.upper())
    _row("IEC Standard", s.standard or "—")
    _row("Module ID", s.module_id or "—")
    _row("Started", s.started_at.isoformat() if s.started_at else "—")
    _row("Ended", s.ended_at.isoformat() if s.ended_at else "—")
    _row("Status", s.status)
    _row("Verdict", verdict)
    _row("Generated", datetime.utcnow().isoformat() + "Z")

    doc.add_paragraph()
    h = doc.add_paragraph()
    h.add_run("Measurement Trace").bold = True

    img_stream = io.BytesIO(chart_png)
    doc.add_picture(img_stream, width=Inches(6.2))

    doc.add_paragraph()
    h2 = doc.add_paragraph()
    h2.add_run("Raw Data (sampled)").bold = True

    table = doc.add_table(rows=1, cols=5)
    table.style = "Light Grid Accent 1"
    hdr = table.rows[0].cells
    hdr[0].text = "Timestamp (UTC)"
    hdr[1].text = "V (V)"
    hdr[2].text = "I (A)"
    hdr[3].text = "P (W)"
    hdr[4].text = "Step"

    # Sample up to ~120 rows so the report stays readable.
    stride = max(1, len(rows) // 120) if rows else 1
    for r in rows[::stride]:
        cells = table.add_row().cells
        cells[0].text = r.ts.isoformat()
        cells[1].text = f"{r.v:.4f}" if r.v is not None else ""
        cells[2].text = f"{r.i:.4f}" if r.i is not None else ""
        cells[3].text = f"{r.p:.4f}" if r.p is not None else ""
        cells[4].text = str(r.step or 0)

    fname = f"{s.test_id}_{s.id[:8]}_{datetime.utcnow().strftime('%Y%m%dT%H%M%S')}.docx"
    fpath = os.path.join(REPORTS_DIR, fname)
    doc.save(fpath)

    db.add(Report(session_id=s.id, format="word", filepath=fpath))
    db.commit()
    return fpath


# ----- PDF report ------------------------------------------------------------

def generate_pdf(db: Session, session_id: str) -> str:
    s, rows = _load(db, session_id)
    chart_png = _make_chart_png(rows)
    verdict, verdict_color = _verdict_and_color(
        json.loads(s.result_json) if s.result_json else None
    )

    fname = f"{s.test_id}_{s.id[:8]}_{datetime.utcnow().strftime('%Y%m%dT%H%M%S')}.pdf"
    fpath = os.path.join(REPORTS_DIR, fname)

    pdf = SimpleDocTemplate(
        fpath,
        pagesize=A4,
        leftMargin=1.8 * cm,
        rightMargin=1.8 * cm,
        topMargin=1.6 * cm,
        bottomMargin=1.6 * cm,
        title=f"{PRODUCT_NAME} Report {s.id[:8]}",
        author=BRAND_NAME,
    )
    styles = getSampleStyleSheet()
    brand_style = ParagraphStyle(
        "brand",
        parent=styles["Title"],
        textColor=colors.HexColor(BRAND_PRIMARY),
        alignment=1,
        fontSize=20,
    )
    tagline_style = ParagraphStyle(
        "tagline", parent=styles["Italic"], alignment=1, fontSize=10, textColor=colors.grey
    )
    verdict_style = ParagraphStyle(
        "verdict",
        parent=styles["Heading2"],
        textColor=colors.HexColor(verdict_color),
        alignment=1,
    )

    story: list = []
    story.append(Paragraph(f"{BRAND_NAME} — {PRODUCT_NAME}", brand_style))
    story.append(Paragraph(BRAND_TAGLINE, tagline_style))
    story.append(Spacer(1, 0.5 * cm))
    story.append(Paragraph(f"Verdict: {verdict}", verdict_style))
    story.append(Spacer(1, 0.4 * cm))

    meta_rows = [
        ["Session ID", s.id],
        ["Test", s.test_id.upper()],
        ["IEC Standard", s.standard or "—"],
        ["Module ID", s.module_id or "—"],
        ["Started", s.started_at.isoformat() if s.started_at else "—"],
        ["Ended", s.ended_at.isoformat() if s.ended_at else "—"],
        ["Status", s.status],
        ["Generated", datetime.utcnow().isoformat() + "Z"],
    ]
    meta_tbl = Table(meta_rows, hAlign="LEFT", colWidths=[4.2 * cm, 11 * cm])
    meta_tbl.setStyle(
        TableStyle(
            [
                ("BACKGROUND", (0, 0), (0, -1), colors.HexColor("#E8F1EE")),
                ("FONTNAME", (0, 0), (0, -1), "Helvetica-Bold"),
                ("GRID", (0, 0), (-1, -1), 0.25, colors.grey),
                ("FONTSIZE", (0, 0), (-1, -1), 9),
                ("VALIGN", (0, 0), (-1, -1), "MIDDLE"),
            ]
        )
    )
    story.append(meta_tbl)
    story.append(Spacer(1, 0.5 * cm))

    story.append(Paragraph("Measurement Trace", styles["Heading3"]))
    chart_img = RLImage(io.BytesIO(chart_png), width=16 * cm, height=7.5 * cm)
    story.append(chart_img)
    story.append(Spacer(1, 0.5 * cm))

    story.append(Paragraph("Raw Data (sampled)", styles["Heading3"]))
    table_data = [["Timestamp (UTC)", "V (V)", "I (A)", "P (W)", "Step"]]
    stride = max(1, len(rows) // 60) if rows else 1
    for r in rows[::stride]:
        table_data.append(
            [
                r.ts.isoformat(timespec="seconds"),
                f"{r.v:.4f}" if r.v is not None else "",
                f"{r.i:.4f}" if r.i is not None else "",
                f"{r.p:.4f}" if r.p is not None else "",
                str(r.step or 0),
            ]
        )
    data_tbl = Table(table_data, hAlign="LEFT", repeatRows=1)
    data_tbl.setStyle(
        TableStyle(
            [
                ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor(BRAND_PRIMARY)),
                ("TEXTCOLOR", (0, 0), (-1, 0), colors.white),
                ("FONTNAME", (0, 0), (-1, 0), "Helvetica-Bold"),
                ("FONTSIZE", (0, 0), (-1, -1), 8),
                ("GRID", (0, 0), (-1, -1), 0.2, colors.lightgrey),
                ("ROWBACKGROUNDS", (0, 1), (-1, -1), [colors.white, colors.HexColor("#F6F8F7")]),
            ]
        )
    )
    story.append(data_tbl)

    pdf.build(story)

    db.add(Report(session_id=s.id, format="pdf", filepath=fpath))
    db.commit()
    return fpath
