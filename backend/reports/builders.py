"""PDF (reportlab) and DOCX (python-docx) builders sharing a section registry.

Both builders walk the same ordered list of section IDs from the request and
emit the corresponding block. Unknown IDs are silently ignored so the
front-end and back-end can evolve independently.
"""
from __future__ import annotations

import base64
import io
import math
from datetime import datetime, timezone
from typing import Callable

from .registry import (
    ALL_GRAPHS,
    ALL_SECTIONS,
    ALL_TABLES,
    GRAPH_LABELS,
    Reading,
    ReportRequest,
    TABLE_LABELS,
    normalize_sections,
)


# --------------------------------------------------------------------------
# Shared helpers
# --------------------------------------------------------------------------
def _graph_series(readings: list[Reading], key: str) -> list[tuple[float, float]]:
    out: list[tuple[float, float]] = []
    for r in readings:
        if key == "vf_vs_t":
            if r.vf is None or r.temperature is None:
                continue
            out.append((r.temperature, r.vf))
            continue
        y = getattr(r, _GRAPH_ATTR[key], None)
        if y is None:
            continue
        out.append((r.timestamp, y))
    return out


_GRAPH_ATTR: dict[str, str] = {
    "voltage": "voltage",
    "current": "current",
    "power": "power",
    "temperature": "temperature",
    "rh": "rh",
    "tj": "tj",
}


def _summary_rows(readings: list[Reading]) -> list[list[str]]:
    if not readings:
        return [["No readings", "—", "—", "—"]]
    rows: list[list[str]] = []
    for label, attr in (
        ("Voltage (V)", "voltage"),
        ("Current (A)", "current"),
        ("Power (W)", "power"),
        ("Temperature (°C)", "temperature"),
        ("RH (%)", "rh"),
    ):
        vals = [getattr(r, attr) for r in readings if getattr(r, attr) is not None]
        if not vals:
            continue
        avg = sum(vals) / len(vals)
        rows.append([
            label,
            f"{min(vals):.4f}",
            f"{avg:.4f}",
            f"{max(vals):.4f}",
        ])
    return rows or [["No numeric series", "—", "—", "—"]]


def _decimate(readings: list[Reading], target: int = 20) -> list[Reading]:
    if len(readings) <= target:
        return readings
    step = max(1, len(readings) // target)
    return readings[::step][:target]


# --------------------------------------------------------------------------
# PDF builder (reportlab)
# --------------------------------------------------------------------------
def build_pdf(req: ReportRequest) -> bytes:
    from reportlab.lib import colors
    from reportlab.lib.pagesizes import A4
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.lib.units import mm
    from reportlab.platypus import (
        SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle, PageBreak, Image,
    )

    sections = normalize_sections(req.sections, ALL_SECTIONS)
    graphs = normalize_sections(req.graphs, ALL_GRAPHS)
    tables = normalize_sections(req.tables, ALL_TABLES)

    buf = io.BytesIO()
    doc = SimpleDocTemplate(
        buf,
        pagesize=A4,
        leftMargin=18 * mm,
        rightMargin=18 * mm,
        topMargin=18 * mm,
        bottomMargin=22 * mm,
        title=f"Agnipariksha Report — {req.effective_test_name}",
    )
    styles = getSampleStyleSheet()
    h1 = ParagraphStyle("h1", parent=styles["Heading1"], textColor=colors.HexColor("#111827"))
    h2 = ParagraphStyle("h2", parent=styles["Heading2"], textColor=colors.HexColor("#111827"))
    body = styles["BodyText"]
    muted = ParagraphStyle("muted", parent=body, textColor=colors.HexColor("#6b7280"), fontSize=9)

    flow: list = []

    def _hr() -> None:
        flow.append(Spacer(1, 4))

    handlers: dict[str, Callable[[], None]] = {
        "header": lambda: _pdf_header(req, flow, h1, muted),
        "test_description": lambda: _pdf_test_description(req, flow, h2, body),
        "iec_clause": lambda: _pdf_iec_clause(req, flow, h2, body),
        "parameters": lambda: _pdf_parameters(req, flow, h2),
        "graphs": lambda: _pdf_graphs(req, graphs, flow, h2, muted),
        "tables": lambda: _pdf_tables(req, tables, flow, h2),
        "pass_fail": lambda: _pdf_pass_fail(req, flow, h2),
        "raw_data_path": lambda: _pdf_raw_path(req, flow, h2, body),
        "error_log": lambda: _pdf_error_log(req, flow, h2, body),
        "troubleshooting": lambda: _pdf_troubleshooting(req, flow, h2, body),
        "signature": lambda: _pdf_signature(req, flow, h2),
        "photos": lambda: _pdf_photos(req, flow, h2, muted),
    }

    for sid in sections:
        fn = handlers.get(sid)
        if fn is None:
            continue
        fn()
        _hr()

    qr_url = _qr_target(req)

    def _on_page(canvas, _doc) -> None:
        canvas.saveState()
        canvas.setFont("Helvetica", 8)
        canvas.setFillColor(colors.HexColor("#6b7280"))
        canvas.drawString(
            18 * mm, 10 * mm,
            f"Agnipariksha · {req.lab_name or 'ASA PV Lab'} · "
            f"Page {canvas.getPageNumber()}",
        )
        canvas.drawRightString(
            A4[0] - 18 * mm, 10 * mm,
            f"Run {req.effective_run_id}",
        )
        try:
            from reportlab.graphics.barcode import qr
            from reportlab.graphics.shapes import Drawing
            from reportlab.graphics import renderPDF
            code = qr.QrCodeWidget(qr_url)
            bounds = code.getBounds()
            w = bounds[2] - bounds[0]
            h = bounds[3] - bounds[1]
            size = 18 * mm
            d = Drawing(size, size, transform=[size / w, 0, 0, size / h, 0, 0])
            d.add(code)
            renderPDF.draw(d, canvas, (A4[0] - 18 * mm - size) / 2, 4 * mm)
        except Exception:
            # QR is best-effort; never fail the report over it
            pass
        canvas.restoreState()

    doc.build(flow, onFirstPage=_on_page, onLaterPages=_on_page)
    return buf.getvalue()


def _qr_target(req: ReportRequest) -> str:
    base = (req.qr_base_url or "").rstrip("/")
    return f"{base}/runs/{req.effective_run_id}" if base else f"/runs/{req.effective_run_id}"


def _pdf_header(req, flow, h1, muted) -> None:
    from reportlab.platypus import Paragraph, Spacer
    flow.append(Paragraph("AGNIPARIKSHA", h1))
    flow.append(Paragraph("PV Module Reliability Test Report", muted))
    if req.lab_name:
        flow.append(Paragraph(req.lab_name, muted))
    flow.append(Spacer(1, 6))
    flow.append(Paragraph(req.effective_test_name, h1))


def _pdf_test_description(req, flow, h2, body) -> None:
    from reportlab.platypus import Paragraph
    flow.append(Paragraph("Test Description", h2))
    desc = req.notes or (
        f"Telemetry capture for {req.effective_test_name} per "
        f"{req.standard or 'IEC 61215-2'}."
    )
    flow.append(Paragraph(desc, body))


def _pdf_iec_clause(req, flow, h2, body) -> None:
    from reportlab.platypus import Paragraph
    flow.append(Paragraph("IEC Clause", h2))
    flow.append(Paragraph(
        f"Standard: {req.standard or '—'} · Clause: {req.iec_clause or '—'}",
        body,
    ))


def _pdf_parameters(req, flow, h2) -> None:
    from reportlab.lib import colors
    from reportlab.platypus import Paragraph, Table, TableStyle
    flow.append(Paragraph("Parameters", h2))
    started = datetime.now(timezone.utc).isoformat(timespec="seconds")
    rows = [
        ["Parameter", "Value"],
        ["Run ID", req.effective_run_id],
        ["Test Type", req.effective_test_name],
        ["Standard", req.standard or "—"],
        ["IEC Clause", req.iec_clause or "—"],
        ["Module ID", req.module_id or "—"],
        ["Operator", req.operator or "—"],
        ["Generated", started],
        ["Total readings", str(len(req.readings))],
    ]
    t = Table(rows, colWidths=[60 * 1.5, 360])
    t.setStyle(TableStyle([
        ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#111827")),
        ("TEXTCOLOR", (0, 0), (-1, 0), colors.HexColor("#ffa500")),
        ("FONTNAME", (0, 0), (-1, 0), "Helvetica-Bold"),
        ("GRID", (0, 0), (-1, -1), 0.25, colors.HexColor("#9ca3af")),
        ("ROWBACKGROUNDS", (0, 1), (-1, -1), [colors.white, colors.HexColor("#f5f5f5")]),
        ("FONTSIZE", (0, 0), (-1, -1), 9),
    ]))
    flow.append(t)


def _pdf_graphs(req, graphs, flow, h2, muted) -> None:
    from reportlab.platypus import Paragraph, Spacer
    flow.append(Paragraph("Graphs", h2))
    if not graphs:
        flow.append(Paragraph("No graphs selected.", muted))
        return
    for gid in graphs:
        img = _render_graph_png(req.readings, gid)
        flow.append(Paragraph(GRAPH_LABELS.get(gid, gid), muted))
        if img is None:
            flow.append(Paragraph("(no data)", muted))
        else:
            from reportlab.platypus import Image
            flow.append(Image(io.BytesIO(img), width=420, height=140))
        flow.append(Spacer(1, 6))


def _pdf_tables(req, tables, flow, h2) -> None:
    from reportlab.lib import colors
    from reportlab.platypus import Paragraph, Table, TableStyle, Spacer
    flow.append(Paragraph("Tables", h2))
    for tid in tables:
        flow.append(Paragraph(TABLE_LABELS.get(tid, tid), h2))
        rows = _table_rows(req.readings, tid)
        if not rows:
            continue
        t = Table(rows, repeatRows=1)
        t.setStyle(TableStyle([
            ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#111827")),
            ("TEXTCOLOR", (0, 0), (-1, 0), colors.HexColor("#ffa500")),
            ("FONTNAME", (0, 0), (-1, 0), "Helvetica-Bold"),
            ("GRID", (0, 0), (-1, -1), 0.25, colors.HexColor("#9ca3af")),
            ("FONTSIZE", (0, 0), (-1, -1), 7),
        ]))
        flow.append(t)
        flow.append(Spacer(1, 6))


def _table_rows(readings: list[Reading], tid: str) -> list[list[str]]:
    if tid == "summary":
        return [["Series", "Min", "Avg", "Max"], *_summary_rows(readings)]
    src = readings if tid == "raw" else _decimate(readings)
    head = ["t", "V", "I", "P", "T", "RH", "Tj", "Vf"]
    out: list[list[str]] = [head]
    for r in src:
        out.append([
            f"{r.timestamp:.2f}",
            "" if r.voltage is None else f"{r.voltage:.3f}",
            "" if r.current is None else f"{r.current:.3f}",
            "" if r.power is None else f"{r.power:.3f}",
            "" if r.temperature is None else f"{r.temperature:.2f}",
            "" if r.rh is None else f"{r.rh:.1f}",
            "" if r.tj is None else f"{r.tj:.2f}",
            "" if r.vf is None else f"{r.vf:.3f}",
        ])
    return out


def _pdf_pass_fail(req, flow, h2) -> None:
    from reportlab.lib import colors
    from reportlab.platypus import Paragraph, Table, TableStyle
    flow.append(Paragraph("Pass / Fail", h2))
    verdict = (req.result or "IN PROGRESS").upper()
    delta_str = "—" if req.delta_pmax_percent is None else f"{req.delta_pmax_percent:.2f}%"
    pre = "—" if req.pre_max_power is None else f"{req.pre_max_power:.3f}"
    post = "—" if req.post_max_power is None else f"{req.post_max_power:.3f}"
    thr = "—" if req.threshold_percent is None else f"{req.threshold_percent:.2f}%"
    rows = [
        ["Pre Pmax", "Post Pmax", "ΔPmax", "Threshold", "Verdict"],
        [pre, post, delta_str, thr, verdict],
    ]
    t = Table(rows)
    is_pass = verdict == "PASS"
    t.setStyle(TableStyle([
        ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#111827")),
        ("TEXTCOLOR", (0, 0), (-1, 0), colors.HexColor("#ffa500")),
        ("FONTNAME", (0, 0), (-1, 0), "Helvetica-Bold"),
        ("GRID", (0, 0), (-1, -1), 0.25, colors.HexColor("#9ca3af")),
        ("TEXTCOLOR", (-1, 1), (-1, 1),
         colors.HexColor("#15803d") if is_pass else colors.HexColor("#b91c1c")),
        ("FONTNAME", (-1, 1), (-1, 1), "Helvetica-Bold"),
    ]))
    flow.append(t)


def _pdf_raw_path(req, flow, h2, body) -> None:
    from reportlab.platypus import Paragraph
    flow.append(Paragraph("Raw Data Path", h2))
    flow.append(Paragraph(req.raw_data_path or "(not specified)", body))


def _pdf_error_log(req, flow, h2, body) -> None:
    from reportlab.platypus import Paragraph
    flow.append(Paragraph("Error Log", h2))
    if not req.error_log:
        flow.append(Paragraph("No errors recorded.", body))
        return
    for line in req.error_log:
        flow.append(Paragraph(f"• {line}", body))


def _pdf_troubleshooting(req, flow, h2, body) -> None:
    from reportlab.platypus import Paragraph
    flow.append(Paragraph("Troubleshooting", h2))
    if not req.troubleshooting:
        flow.append(Paragraph(
            "No troubleshooting notes. Refer to the test SOP for diagnostic steps.",
            body,
        ))
        return
    for line in req.troubleshooting:
        flow.append(Paragraph(f"• {line}", body))


def _pdf_signature(req, flow, h2) -> None:
    from reportlab.lib import colors
    from reportlab.platypus import Paragraph, Table, TableStyle
    flow.append(Paragraph("Signature", h2))
    rows = [
        ["Operator", req.operator or "____________________"],
        ["Date", datetime.now(timezone.utc).strftime("%Y-%m-%d")],
        ["Signature", "____________________"],
    ]
    t = Table(rows, colWidths=[100, 360])
    t.setStyle(TableStyle([
        ("GRID", (0, 0), (-1, -1), 0.25, colors.HexColor("#9ca3af")),
        ("FONTSIZE", (0, 0), (-1, -1), 9),
    ]))
    flow.append(t)


def _pdf_photos(req, flow, h2, muted) -> None:
    from reportlab.platypus import Paragraph, Image, Spacer
    flow.append(Paragraph("Photos", h2))
    if not req.photos:
        flow.append(Paragraph("No photos attached.", muted))
        return
    for idx, b64 in enumerate(req.photos, start=1):
        try:
            raw = base64.b64decode(b64.split(",", 1)[-1])
            flow.append(Paragraph(f"Photo {idx}", muted))
            flow.append(Image(io.BytesIO(raw), width=300, height=200))
            flow.append(Spacer(1, 4))
        except Exception:
            flow.append(Paragraph(f"Photo {idx}: (invalid image data)", muted))


# --------------------------------------------------------------------------
# Graph rendering — matplotlib backend, no display.
# --------------------------------------------------------------------------
def _render_graph_png(readings: list[Reading], gid: str) -> bytes | None:
    series = _graph_series(readings, gid)
    if not series:
        return None
    try:
        import matplotlib
        matplotlib.use("Agg")
        import matplotlib.pyplot as plt
    except Exception:
        return None
    xs = [p[0] for p in series]
    ys = [p[1] for p in series]
    fig, ax = plt.subplots(figsize=(7.2, 2.4), dpi=120)
    ax.plot(xs, ys, linewidth=1.2)
    ax.set_title(GRAPH_LABELS.get(gid, gid), fontsize=9)
    ax.grid(True, linewidth=0.3, alpha=0.4)
    ax.tick_params(labelsize=7)
    fig.tight_layout()
    buf = io.BytesIO()
    fig.savefig(buf, format="png")
    plt.close(fig)
    return buf.getvalue()


# --------------------------------------------------------------------------
# DOCX builder (python-docx)
# --------------------------------------------------------------------------
def build_docx(req: ReportRequest) -> bytes:
    from docx import Document
    from docx.shared import Inches, Pt, RGBColor

    sections = normalize_sections(req.sections, ALL_SECTIONS)
    graphs = normalize_sections(req.graphs, ALL_GRAPHS)
    tables = normalize_sections(req.tables, ALL_TABLES)

    doc = Document()

    def _h2(text: str) -> None:
        doc.add_heading(text, level=2)

    def _muted(text: str) -> None:
        p = doc.add_paragraph()
        run = p.add_run(text)
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(0x6B, 0x72, 0x80)

    handlers: dict[str, Callable[[], None]] = {
        "header": lambda: _docx_header(req, doc),
        "test_description": lambda: (_h2("Test Description"), doc.add_paragraph(
            req.notes or f"Telemetry capture for {req.effective_test_name} per "
            f"{req.standard or 'IEC 61215-2'}."
        )),
        "iec_clause": lambda: (_h2("IEC Clause"), doc.add_paragraph(
            f"Standard: {req.standard or '—'} · Clause: {req.iec_clause or '—'}"
        )),
        "parameters": lambda: _docx_parameters(req, doc),
        "graphs": lambda: _docx_graphs(req, graphs, doc, _muted),
        "tables": lambda: _docx_tables(req, tables, doc),
        "pass_fail": lambda: _docx_pass_fail(req, doc),
        "raw_data_path": lambda: (_h2("Raw Data Path"), doc.add_paragraph(
            req.raw_data_path or "(not specified)"
        )),
        "error_log": lambda: _docx_error_log(req, doc),
        "troubleshooting": lambda: _docx_troubleshooting(req, doc),
        "signature": lambda: _docx_signature(req, doc),
        "photos": lambda: _docx_photos(req, doc, _muted),
    }

    for sid in sections:
        fn = handlers.get(sid)
        if fn is None:
            continue
        fn()

    # Footer with run id + QR target URL (DOCX has no native QR rendering)
    section = doc.sections[0]
    footer = section.footer.paragraphs[0]
    footer.text = (
        f"Agnipariksha · {req.lab_name or 'ASA PV Lab'} · "
        f"Run {req.effective_run_id} · {_qr_target(req)}"
    )

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def _docx_header(req, doc) -> None:
    from docx.shared import Pt, RGBColor
    title = doc.add_heading("AGNIPARIKSHA", level=0)
    for run in title.runs:
        run.font.color.rgb = RGBColor(0xFF, 0xA5, 0x00)
    sub = doc.add_paragraph().add_run("PV Module Reliability Test Report")
    sub.italic = True
    sub.font.size = Pt(10)
    if req.lab_name:
        doc.add_paragraph(req.lab_name)
    doc.add_heading(req.effective_test_name, level=1)


def _docx_parameters(req, doc) -> None:
    doc.add_heading("Parameters", level=2)
    rows = [
        ("Run ID", req.effective_run_id),
        ("Test Type", req.effective_test_name),
        ("Standard", req.standard or "—"),
        ("IEC Clause", req.iec_clause or "—"),
        ("Module ID", req.module_id or "—"),
        ("Operator", req.operator or "—"),
        ("Generated", datetime.now(timezone.utc).isoformat(timespec="seconds")),
        ("Total readings", str(len(req.readings))),
    ]
    t = doc.add_table(rows=len(rows), cols=2)
    t.style = "Light List Accent 1"
    for i, (k, v) in enumerate(rows):
        t.rows[i].cells[0].text = k
        t.rows[i].cells[1].text = v


def _docx_graphs(req, graphs, doc, muted) -> None:
    from docx.shared import Inches
    doc.add_heading("Graphs", level=2)
    if not graphs:
        muted("No graphs selected.")
        return
    for gid in graphs:
        doc.add_paragraph(GRAPH_LABELS.get(gid, gid))
        img = _render_graph_png(req.readings, gid)
        if img is None:
            muted("(no data)")
            continue
        doc.add_picture(io.BytesIO(img), width=Inches(6.0))


def _docx_tables(req, tables, doc) -> None:
    doc.add_heading("Tables", level=2)
    for tid in tables:
        doc.add_heading(TABLE_LABELS.get(tid, tid), level=3)
        rows = _table_rows(req.readings, tid)
        if not rows:
            continue
        t = doc.add_table(rows=len(rows), cols=len(rows[0]))
        t.style = "Light Grid Accent 1"
        for r_idx, row in enumerate(rows):
            for c_idx, cell in enumerate(row):
                t.rows[r_idx].cells[c_idx].text = cell


def _docx_pass_fail(req, doc) -> None:
    from docx.shared import RGBColor
    doc.add_heading("Pass / Fail", level=2)
    verdict = (req.result or "IN PROGRESS").upper()
    p = doc.add_paragraph()
    run = p.add_run(verdict)
    run.bold = True
    run.font.color.rgb = (
        RGBColor(0x15, 0x80, 0x3D) if verdict == "PASS"
        else RGBColor(0xB9, 0x1C, 0x1C) if verdict == "FAIL"
        else RGBColor(0x6B, 0x72, 0x80)
    )
    rows = [
        ("Pre Pmax", "—" if req.pre_max_power is None else f"{req.pre_max_power:.3f}"),
        ("Post Pmax", "—" if req.post_max_power is None else f"{req.post_max_power:.3f}"),
        ("ΔPmax",
         "—" if req.delta_pmax_percent is None else f"{req.delta_pmax_percent:.2f}%"),
        ("Threshold",
         "—" if req.threshold_percent is None else f"{req.threshold_percent:.2f}%"),
    ]
    t = doc.add_table(rows=len(rows), cols=2)
    t.style = "Light List Accent 1"
    for i, (k, v) in enumerate(rows):
        t.rows[i].cells[0].text = k
        t.rows[i].cells[1].text = v


def _docx_error_log(req, doc) -> None:
    doc.add_heading("Error Log", level=2)
    if not req.error_log:
        doc.add_paragraph("No errors recorded.")
        return
    for line in req.error_log:
        doc.add_paragraph(line, style="List Bullet")


def _docx_troubleshooting(req, doc) -> None:
    doc.add_heading("Troubleshooting", level=2)
    if not req.troubleshooting:
        doc.add_paragraph(
            "No troubleshooting notes. Refer to the test SOP for diagnostic steps."
        )
        return
    for line in req.troubleshooting:
        doc.add_paragraph(line, style="List Bullet")


def _docx_signature(req, doc) -> None:
    doc.add_heading("Signature", level=2)
    rows = [
        ("Operator", req.operator or "____________________"),
        ("Date", datetime.now(timezone.utc).strftime("%Y-%m-%d")),
        ("Signature", "____________________"),
    ]
    t = doc.add_table(rows=len(rows), cols=2)
    t.style = "Light List Accent 1"
    for i, (k, v) in enumerate(rows):
        t.rows[i].cells[0].text = k
        t.rows[i].cells[1].text = v


def _docx_photos(req, doc, muted) -> None:
    from docx.shared import Inches
    doc.add_heading("Photos", level=2)
    if not req.photos:
        muted("No photos attached.")
        return
    for idx, b64 in enumerate(req.photos, start=1):
        try:
            raw = base64.b64decode(b64.split(",", 1)[-1])
            doc.add_paragraph(f"Photo {idx}")
            doc.add_picture(io.BytesIO(raw), width=Inches(4.0))
        except Exception:
            muted(f"Photo {idx}: (invalid image data)")
